from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


TOPICS = {
    "Agent/PEAS": ["PEAS", "Agent", "agent", "智能体", "性能度量", "感知器", "执行器"],
    "状态空间搜索": ["搜索", "BFS", "DFS", "A*", "A star", "启发", "贪婪", "一致代价", "UCS", "宽度优先", "深度优先"],
    "局部搜索": ["爬山", "模拟退火", "遗传算法", "局部搜索", "退火"],
    "博弈搜索": ["博弈", "minimax", "极大极小", "alpha", "α", "beta", "β", "剪枝"],
    "CSP": ["CSP", "约束满足", "约束传播", "回溯", "八皇后", "数独"],
    "逻辑/归结": ["命题逻辑", "一阶逻辑", "CNF", "归结", "真值表", "谓词", "Skolem", "合取范式"],
    "概率基础": ["概率", "条件概率", "Bayes", "贝叶斯公式", "独立性", "扑克牌"],
    "贝叶斯网络": ["贝叶斯网络", "Bayesian Network", "BN", "条件独立", "联合分布", "枚举推理"],
    "HMM/时间概率": ["HMM", "隐藏马尔可夫", "隐马尔可夫", "马尔可夫", "过滤", "预测", "平滑", "Viterbi", "维特比"],
    "朴素贝叶斯": ["朴素贝叶斯", "Naive Bayes", "naive Bayes", "拉普拉斯", "Laplace"],
    "决策树": ["决策树", "信息熵", "熵", "信息增益", "Gain", "ID3"],
    "机器学习评估": ["训练集", "验证集", "测试集", "K折", "交叉验证", "监督学习", "无监督学习", "半监督", "过拟合", "正则化"],
    "神经网络": ["神经网络", "感知器", "前向传播", "反向传播", "梯度下降", "激活函数", "Softmax"],
    "CNN/ResNet": ["CNN", "卷积", "池化", "卷积核", "ResNet", "残差", "深度学习"],
}


def read_pdf(path: Path, max_pages: int) -> tuple[str, int | str]:
    try:
        reader = PdfReader(str(path))
        text = []
        for page in reader.pages[:max_pages]:
            text.append(page.extract_text() or "")
        return "\n".join(text), len(reader.pages)
    except Exception as exc:
        return "", f"ERR:{exc.__class__.__name__}"


def read_pptx(path: Path) -> tuple[str, int | str]:
    try:
        with zipfile.ZipFile(path) as archive:
            slide_names = [
                name
                for name in archive.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            ]
            slide_names.sort(key=lambda item: int(re.search(r"slide(\d+)\.xml", item).group(1)))
            namespace = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
            slides = []
            for name in slide_names:
                xml_root = ET.fromstring(archive.read(name))
                parts = [node.text.strip() for node in xml_root.iter(namespace + "t") if node.text and node.text.strip()]
                if parts:
                    slides.append(" ".join(parts))
            return "\n".join(slides), len(slide_names)
    except Exception as exc:
        return "", f"ERR:{exc.__class__.__name__}"


def read_docx(path: Path) -> tuple[str, int | str]:
    try:
        doc = Document(str(path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs), len(doc.paragraphs)
    except Exception as exc:
        return "", f"ERR:{exc.__class__.__name__}"


def read_text(path: Path, max_pdf_pages: int) -> tuple[str, int | str | None]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path, max_pdf_pages)
    if suffix == ".pptx":
        return read_pptx(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".md":
        return path.read_text(encoding="utf-8", errors="ignore"), None
    return "", None


def topic_counts(text: str) -> dict[str, int]:
    counts: dict[str, int] = {}
    for topic, patterns in TOPICS.items():
        total = 0
        for pattern in patterns:
            total += len(re.findall(re.escape(pattern), text, flags=re.IGNORECASE))
        if total:
            counts[topic] = total
    return counts


def collect_files(root: Path) -> list[Path]:
    wanted = {".pdf", ".pptx", ".docx", ".md"}
    ignored_parts = {".git", "__pycache__"}
    files = []
    for path in root.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in wanted:
            continue
        if any(part in ignored_parts for part in path.parts):
            continue
        files.append(path)
    return sorted(files)


def main() -> None:
    parser = argparse.ArgumentParser(description="Scan AI review materials and count topic signals.")
    parser.add_argument("--root", default=".", help="Repository root.")
    parser.add_argument("--max-pdf-pages", type=int, default=20, help="Pages to scan per PDF.")
    parser.add_argument("--json", action="store_true", help="Print JSON instead of Markdown.")
    args = parser.parse_args()

    root = Path(args.root).resolve()
    records = []
    aggregate = {topic: 0 for topic in TOPICS}

    for path in collect_files(root):
        text, page_or_slide_count = read_text(path, args.max_pdf_pages)
        counts = topic_counts(text)
        for topic, count in counts.items():
            aggregate[topic] += count
        records.append(
            {
                "path": str(path.relative_to(root)),
                "kind": path.suffix.lower(),
                "pages_or_slides": page_or_slide_count,
                "chars": len(text),
                "topics": counts,
            }
        )

    result = {
        "aggregate": dict(sorted(aggregate.items(), key=lambda item: item[1], reverse=True)),
        "records": sorted(records, key=lambda item: sum(item["topics"].values()), reverse=True),
    }

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    print("# Topic Signals")
    print()
    for topic, count in result["aggregate"].items():
        print(f"- {topic}: {count}")
    print()
    print("# Documents")
    print()
    for record in result["records"]:
        if not record["topics"]:
            continue
        top = ", ".join(f"{topic}:{count}" for topic, count in list(record["topics"].items())[:5])
        print(f"- `{record['path']}` ({record['pages_or_slides']}): {top}")


if __name__ == "__main__":
    main()
